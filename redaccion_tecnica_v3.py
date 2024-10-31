#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import tkinter as tk
from tkinter import filedialog, messagebox
import csv
import math
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import os

# Variables para almacenar los datos
coordenadas = {}
segmentos_colindantes = []
ruta_plantilla = None  # Variable para almacenar la ruta de la plantilla

# Función para calcular la distancia entre dos puntos
def calcular_distancia(punto1, punto2):
    return math.sqrt((punto2[0] - punto1[0])**2 + (punto1[1] - punto2[1])**2)

# Función para cargar el archivo CSV de coordenadas
def cargar_csv_coordenadas():
    filepath = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
    if not filepath:
        return
    
    coordenadas.clear()
    with open(filepath, newline='', encoding='utf-8-sig') as csvfile:
        reader = csv.DictReader(csvfile, delimiter=';')
        
        for row in reader:
            id_secuencial = int(row['Id'].strip())
            id_original = row['Iden'].strip()
            este = float(row['Este'].strip())
            norte = float(row['Norte'].strip())
            coordenadas[id_secuencial] = (id_original, este, norte)
    
    lbl_coordenadas_cargadas.config(text="Coordenadas cargadas")
    mostrar_segmento_formulario()

# Función para agregar un segmento
def agregar_segmento():
    try:
        inicio = int(entry_inicio.get().strip())
        fin = int(entry_fin.get().strip())
        colindante = entry_colindante.get().strip()
        distancia = entry_distancia.get().strip()
        
        if inicio not in coordenadas or fin not in coordenadas:
            raise ValueError("Los puntos deben estar dentro del rango de coordenadas.")
        
        if len(segmentos_colindantes) > 0 and segmentos_colindantes[-1][1] != inicio:
            raise ValueError("Cada segmento debe continuar donde terminó el anterior.")
        
        if distancia:
            distancia = float(distancia)
        else:
            punto_inicio = coordenadas[inicio][1:]
            punto_fin = coordenadas[fin][1:]
            distancia = calcular_distancia(punto_inicio, punto_fin)
        
        segmentos_colindantes.append((inicio, fin, colindante, distancia))
        listbox_segmentos.insert(tk.END, f"{inicio} al {fin} colinda con {colindante}, distancia: {distancia:.2f} m")
    
    except ValueError as e:
        messagebox.showerror("Error", str(e))
    
    entry_inicio.delete(0, tk.END)
    entry_fin.delete(0, tk.END)
    entry_colindante.delete(0, tk.END)
    entry_distancia.delete(0, tk.END)

# Función para generar el mapa con la tabla de segmentos
def generar_mapa(coordenadas, segmentos_colindantes, nombre_imagen=None):
    fig, ax = plt.subplots(figsize=(10, 10))
    ax.set_title('Mapa de Linderos del Predio')
    ax.set_xlabel('Este (m)')
    ax.set_ylabel('Norte (m)')
    
    x_coords = [coordenadas[id_punto][1] for id_punto in coordenadas]
    y_coords = [coordenadas[id_punto][2] for id_punto in coordenadas]
    x_coords.append(x_coords[0])  # Para cerrar el polígono
    y_coords.append(y_coords[0])
    ax.plot(x_coords, y_coords, 'b-', marker='o')

    for id_punto, (id_original, x, y) in coordenadas.items():
        ax.text(x, y, str(id_original), fontsize=9, ha='right', va='bottom', color='blue')

    for inicio, fin, colindante, distancia in segmentos_colindantes:
        mid_x = (coordenadas[inicio][1] + coordenadas[fin][1]) / 2
        mid_y = (coordenadas[inicio][2] + coordenadas[fin][2]) / 2
        texto_segmento = f"Del {coordenadas[inicio][0]} al {coordenadas[fin][0]}, Colindante: {colindante}"
        ax.text(mid_x, mid_y, texto_segmento, fontsize=8, ha='center', va='center', color='red', backgroundcolor='white')

    plt.grid(True)
    
    if nombre_imagen:
        plt.savefig(nombre_imagen, bbox_inches='tight')
    else:
        plt.show()

# Función para seleccionar la plantilla
def seleccionar_plantilla():
    global ruta_plantilla
    ruta_plantilla = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if ruta_plantilla:
        lbl_plantilla.config(text=f"Plantilla seleccionada: {ruta_plantilla}")

# Función para guardar la redacción en un archivo .docx usando la plantilla seleccionada
def guardar_redaccion_como_docx(redaccion, nombre_imagen):
    if not ruta_plantilla or not os.path.exists(ruta_plantilla):
        messagebox.showerror("Error", "No se ha seleccionado una plantilla válida.")
        return
    
    nombre_predio = entry_nombre_predio.get()
    cedula_catastral = entry_cedula_catastral.get()
    folio_matricula = entry_folio_matricula.get()
    municipio = entry_municipio.get()
    departamento = entry_departamento.get()
    
    filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if filepath:
        doc = Document(ruta_plantilla)
        
        # Reemplazar marcadores en el documento
        for para in doc.paragraphs:
            if '[NOMBRE_PREDIO]' in para.text:
                para.text = para.text.replace('[NOMBRE_PREDIO]', nombre_predio)
            if '[CEDULA_CATASTRAL]' in para.text:
                para.text = para.text.replace('[CEDULA_CATASTRAL]', cedula_catastral)
            if '[FOLIO_MATRICULA]' in para.text:
                para.text = para.text.replace('[FOLIO_MATRICULA]', folio_matricula)
            if '[MUNICIPIO]' in para.text:
                para.text = para.text.replace('[MUNICIPIO]', municipio)
            if '[DEPARTAMENTO]' in para.text:
                para.text = para.text.replace('[DEPARTAMENTO]', departamento)
            if '[REDACCION]' in para.text:
                para.text = para.text.replace('[REDACCION]', redaccion)
        
        # Insertar la Tabla 1: Coordenadas de los puntos
        doc.add_heading('Tabla 1: Coordenadas de los Puntos', level=1)
        tabla1 = doc.add_table(rows=1, cols=3)
        hdr_cells = tabla1.rows[0].cells
        hdr_cells[0].text = 'Iden'
        hdr_cells[1].text = 'Coordenada Este'
        hdr_cells[2].text = 'Coordenada Norte'

        for iden, (id_original, este, norte) in coordenadas.items():
            row_cells = tabla1.add_row().cells
            row_cells[0].text = str(id_original)
            row_cells[1].text = f"{este:.3f}"
            row_cells[2].text = f"{norte:.3f}"
        
        doc.add_paragraph("\n")

        # Insertar la Tabla 2: Segmentos del Predio
        doc.add_heading('Tabla 2: Segmentos del Predio', level=1)
        tabla2 = doc.add_table(rows=1, cols=3)
        hdr_cells = tabla2.rows[0].cells
        hdr_cells[0].text = 'Segmento'
        hdr_cells[1].text = 'Colindante'
        hdr_cells[2].text = 'Distancia (m)'

        for inicio, fin, colindante, distancia in segmentos_colindantes:
            row_cells = tabla2.add_row().cells
            row_cells[0].text = f"{coordenadas[inicio][0]} - {coordenadas[fin][0]}"
            row_cells[1].text = colindante
            row_cells[2].text = f"{distancia:.2f}"
        
        for table in [tabla1, tabla2]:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Arial'
        
        doc.add_page_break()
        doc.add_picture(nombre_imagen, width=Inches(6))

        doc.save(filepath)
        messagebox.showinfo("Guardado", "La redacción y las tablas se han guardado correctamente.\n\n"
                                        "Este programa ha sido desarrollado por el Ing. Luis Miguel Guerrero.\n"
                                        "Para consultas o sugerencias, no dude en ponerse en contacto:\n"
                                        "Correo: lmiguelguerrero@outlook.com\n"
                                        "Gracias por utilizar este software. Espero que le sea de gran utilidad en sus proyectos.")

# Función para calcular el azimut entre dos puntos
def calcular_azimut(punto1, punto2):
    delta_este = punto2[0] - punto1[0]
    delta_norte = punto2[1] - punto1[1]
    azimut = math.degrees(math.atan2(delta_este, delta_norte))
    if azimut < 0:
        azimut += 360
    return azimut

# Función para convertir el azimut a rumbo
def convertir_azimut_a_rumbo(azimut):
    if 0 <= azimut < 90:
        rumbo = "noreste"
    elif 90 <= azimut < 180:
        rumbo = "sureste"
    elif 180 <= azimut < 270:
        rumbo = "suroeste"
    elif 270 <= azimut <= 360:
        rumbo = "noroeste"
    else:
        rumbo = "desconocido"
    return rumbo

# Función para generar la redacción
def generar_redaccion(coordenadas, segmentos_colindantes):
    redaccion = ""
    
    for i, segmento in enumerate(segmentos_colindantes, start=1):
        inicio, fin, colindante, distancia = segmento
        distancia = round(distancia, 1) if distancia else round(calcular_distancia(coordenadas[inicio][1:], coordenadas[fin][1:]), 1)
        
        redaccion += f"Lindero {i}\n"
        redaccion += f"Inicia en el punto {coordenadas[inicio][0]} con coordenadas N={round(coordenadas[inicio][2], 2)}m, E={round(coordenadas[inicio][1], 2)}m, "
        
        puntos = list(range(inicio, fin+1)) if inicio < fin else list(range(inicio, len(coordenadas) + 1)) + list(range(1, fin+1))
        
        rumbo_anterior = None
        for j in range(len(puntos) - 1):
            azimut = calcular_azimut(coordenadas[puntos[j]][1:], coordenadas[puntos[j+1]][1:])
            rumbo_actual = convertir_azimut_a_rumbo(azimut)
            
            if rumbo_anterior != rumbo_actual:
                redaccion += f"al punto {coordenadas[puntos[j+1]][0]} con coordenadas N={round(coordenadas[puntos[j+1]][2], 2)}m, E={round(coordenadas[puntos[j+1]][1], 2)}m en sentido {rumbo_actual}, "
            else:
                redaccion += f"al punto {coordenadas[puntos[j+1]][0]} con coordenadas N={round(coordenadas[puntos[j+1]][2], 2)}m, E={round(coordenadas[puntos[j+1]][1], 2)}m, "
            
            rumbo_anterior = rumbo_actual

        redaccion = redaccion.rstrip(", ")
        redaccion += f" en una distancia total de {distancia}m colindando con {colindante}\n\n"
        
    return redaccion

# Función para mostrar el formulario de segmento después de cargar coordenadas
def mostrar_segmento_formulario():
    lbl_segmento_formulario.pack(pady=5)
    frame_formulario.pack(pady=5)
    lbl_segmentos_agregados.pack(pady=5)
    listbox_segmentos.pack(pady=5)
    btn_generar_redaccion.pack(pady=20)

# Función para generar la redacción y mostrarla en la interfaz
def generar_redaccion_interfaz():
    if not segmentos_colindantes:
        messagebox.showerror("Error", "No se han agregado segmentos.")
        return

    redaccion = generar_redaccion(coordenadas, segmentos_colindantes)
    text_redaccion.delete("1.0", tk.END)
    text_redaccion.insert(tk.END, redaccion)

    nombre_imagen = "mapa_predio.png"
    generar_mapa(coordenadas, segmentos_colindantes, nombre_imagen)

    guardar_redaccion_como_docx(redaccion, nombre_imagen)

# Crear la ventana principal
root = tk.Tk()
root.title("Redacción Técnica de Linderos")

# Crear y ubicar los elementos en la ventana
frame_detalles_predio = tk.Frame(root)
frame_detalles_predio.pack(pady=5)

lbl_nombre_predio = tk.Label(frame_detalles_predio, text="Nombre del Predio:")
entry_nombre_predio = tk.Entry(frame_detalles_predio, width=30)

lbl_cedula_catastral = tk.Label(frame_detalles_predio, text="Cédula Catastral:")
entry_cedula_catastral = tk.Entry(frame_detalles_predio, width=30)

lbl_folio_matricula = tk.Label(frame_detalles_predio, text="Folio de Matrícula:")
entry_folio_matricula = tk.Entry(frame_detalles_predio, width=30)

lbl_municipio = tk.Label(frame_detalles_predio, text="Municipio:")
entry_municipio = tk.Entry(frame_detalles_predio, width=30)

lbl_departamento = tk.Label(frame_detalles_predio, text="Departamento:")
entry_departamento = tk.Entry(frame_detalles_predio, width=30)

lbl_nombre_predio.grid(row=1, column=0, padx=5, pady=5, sticky='e')
entry_nombre_predio.grid(row=1, column=1, padx=5, pady=5, sticky='w')

lbl_cedula_catastral.grid(row=2, column=0, padx=5, pady=5, sticky='e')
entry_cedula_catastral.grid(row=2, column=1, padx=5, pady=5, sticky='w')

lbl_folio_matricula.grid(row=3, column=0, padx=5, pady=5, sticky='e')
entry_folio_matricula.grid(row=3, column=1, padx=5, pady=5, sticky='w')

lbl_municipio.grid(row=4, column=0, padx=5, pady=5, sticky='e')
entry_municipio.grid(row=4, column=1, padx=5, pady=5, sticky='w')

lbl_departamento.grid(row=5, column=0, padx=5, pady=5, sticky='e')
entry_departamento.grid(row=5, column=1, padx=5, pady=5, sticky='w')

# Botón para seleccionar la plantilla
btn_seleccionar_plantilla = tk.Button(root, text="Seleccionar Plantilla", command=seleccionar_plantilla)
btn_seleccionar_plantilla.pack(pady=5)

lbl_plantilla = tk.Label(root, text="No se ha seleccionado una plantilla")
lbl_plantilla.pack(pady=5)

btn_cargar_coordenadas = tk.Button(root, text="Cargar CSV de Coordenadas", command=cargar_csv_coordenadas)
btn_cargar_coordenadas.pack(pady=5)

lbl_coordenadas_cargadas = tk.Label(root, text="No se han cargado coordenadas")
lbl_coordenadas_cargadas.pack(pady=5)

lbl_segmento_formulario = tk.Label(root, text="Ingrese los segmentos y colindantes:")
frame_formulario = tk.Frame(root)
lbl_inicio = tk.Label(frame_formulario, text="Inicio")
lbl_fin = tk.Label(frame_formulario, text="Fin")
lbl_colindante = tk.Label(frame_formulario, text="Colindante")
lbl_distancia = tk.Label(frame_formulario, text="Distancia (opcional)")
entry_inicio = tk.Entry(frame_formulario)
entry_fin = tk.Entry(frame_formulario)
entry_colindante = tk.Entry(frame_formulario)
entry_distancia = tk.Entry(frame_formulario)
btn_agregar_segmento = tk.Button(frame_formulario, text="Agregar Segmento", command=agregar_segmento)

lbl_inicio.grid(row=0, column=0)
entry_inicio.grid(row=0, column=1)
lbl_fin.grid(row=1, column=0)
entry_fin.grid(row=1, column=1)
lbl_colindante.grid(row=2, column=0)
entry_colindante.grid(row=2, column=1)
lbl_distancia.grid(row=3, column=0)
entry_distancia.grid(row=3, column=1)
btn_agregar_segmento.grid(row=4, columnspan=2, pady=5)

lbl_segmentos_agregados = tk.Label(root, text="Segmentos agregados:")
listbox_segmentos = tk.Listbox(root, width=50, height=10)

btn_generar_redaccion = tk.Button(root, text="Generar Redacción", command=generar_redaccion_interfaz)

text_redaccion = tk.Text(root, height=20, width=100)

root.mainloop()

